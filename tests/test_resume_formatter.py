"""Tests for the resume_formatter module.

Covers:
  - StyleProfile serialisation / defaults
  - ResumeElement construction
  - parse_resume_text  – element classification
  - split_inline_bold  – inline markup splitting
  - apply_original_bold – per-line Jaccard bold application
  - _apply_phrase_bold  – phrase wrapping helper
  - _is_contact_info_line – contact regex
  - _extract_bold_phrases_for_line / _extract_bold_phrases_per_line
  - generate_formatted_pdf  – smoke test (returns valid PDF bytes)
  - generate_formatted_docx – smoke test (returns valid DOCX bytes)
"""

from __future__ import annotations

import io

import pytest

from career_copilot.resume_formatter.structure_parser import (
    ResumeElement,
    StyleProfile,
    _apply_phrase_bold,
    _extract_bold_phrases_for_line,
    _extract_bold_phrases_per_line,
    _is_contact_info_line,
    apply_original_bold,
    parse_resume_text,
    split_inline_bold,
)

# ---------------------------------------------------------------------------
# StyleProfile
# ---------------------------------------------------------------------------


class TestStyleProfile:
    def test_default_values(self):
        p = StyleProfile()
        assert p.name_font_size == 18.0
        assert p.header_font_size == 13.0
        assert p.body_font_size == 11.0
        assert p.base_font_family == "Helvetica"
        assert p.name_align == "left"
        assert p.name_color == "#000000"
        assert p.header_color == "#000000"
        assert p.body_color == "#000000"
        assert p.bullet_char == "•"
        assert p.bold_phrases == []
        assert p.has_section_rule is False
        assert p.section_rule_style == "single"

    def test_default_factory(self):
        assert StyleProfile.default() == StyleProfile()

    def test_to_json_round_trip(self):
        p = StyleProfile(
            name_font_size=20.0,
            header_color="#1f497d",
            bold_phrases=[["Teck Resources", "Teck Resources"]],
            has_section_rule=True,
            section_rule_style="thickThinSmallGap",
        )
        restored = StyleProfile.from_json(p.to_json())
        assert restored == p

    def test_from_json_partial_overrides_defaults(self):
        import json

        data = {"name_font_size": 24.0}
        p = StyleProfile.from_json(json.dumps({**StyleProfile().__dict__, **data}))
        assert p.name_font_size == 24.0
        assert p.body_font_size == 11.0  # default preserved


# ---------------------------------------------------------------------------
# ResumeElement
# ---------------------------------------------------------------------------


class TestResumeElement:
    def test_creation(self):
        el = ResumeElement("name", "Jane Doe")
        assert el.kind == "name"
        assert el.text == "Jane Doe"

    def test_blank_element(self):
        el = ResumeElement("blank", "")
        assert el.kind == "blank"
        assert el.text == ""


# ---------------------------------------------------------------------------
# _is_contact_info_line
# ---------------------------------------------------------------------------


class TestIsContactInfoLine:
    @pytest.mark.parametrize(
        "line",
        [
            "jane@example.com",
            "https://linkedin.com/in/jane",
            "www.janedoe.com",
            "linkedin.com/in/jane",
            "github.com/jane",
            "janedoe.com",
            "416-527-2903",
            "416.527.2903",
            "(416) 527-2903",
            "Vancouver, BC  416 527 2903",
        ],
    )
    def test_contact_lines_detected(self, line):
        assert _is_contact_info_line(line)

    @pytest.mark.parametrize(
        "line",
        [
            "Data Scientist and Process Engineer",
            "Advanced Analytics and Optimization",
            "Vancouver, BC",
            "Results-driven senior engineer",
            "WORK EXPERIENCE",
        ],
    )
    def test_non_contact_lines_rejected(self, line):
        assert not _is_contact_info_line(line)


# ---------------------------------------------------------------------------
# split_inline_bold
# ---------------------------------------------------------------------------


class TestSplitInlineBold:
    def test_no_markers(self):
        assert split_inline_bold("plain text") == [("plain text", False)]

    def test_single_bold_segment(self):
        result = split_inline_bold("Hello **World**")
        assert result == [("Hello ", False), ("World", True)]

    def test_multiple_bold_segments(self):
        result = split_inline_bold("**Acme Corp** as **Engineer**")
        assert result == [("Acme Corp", True), (" as ", False), ("Engineer", True)]

    def test_bold_at_start_and_end(self):
        result = split_inline_bold("**start** middle **end**")
        assert ("start", True) in result
        assert ("end", True) in result
        assert (" middle ", False) in result

    def test_empty_string(self):
        assert split_inline_bold("") == []

    def test_only_bold(self):
        assert split_inline_bold("**only**") == [("only", True)]


# ---------------------------------------------------------------------------
# _apply_phrase_bold
# ---------------------------------------------------------------------------


class TestApplyPhraseBold:
    def test_wraps_matching_phrase(self):
        assert (
            _apply_phrase_bold("Teck Resources, Vancouver", "Teck Resources")
            == "**Teck Resources**, Vancouver"
        )

    def test_case_insensitive(self):
        result = _apply_phrase_bold("teck resources, Vancouver", "Teck Resources")
        assert "**teck resources**" in result

    def test_does_not_double_wrap(self):
        line = "**Teck Resources**, Vancouver"
        result = _apply_phrase_bold(line, "Teck Resources")
        assert result.count("**") == 2

    def test_no_match_returns_unchanged(self):
        line = "Some other company, Vancouver"
        assert _apply_phrase_bold(line, "Teck Resources") == line

    def test_word_boundary_no_partial_match(self):
        # "Dat" should not match inside "Data"
        assert _apply_phrase_bold("Data Science", "Dat") == "Data Science"

    def test_phrase_with_parentheses(self):
        line = "Inventory Management (BOIM) project"
        result = _apply_phrase_bold(line, "Inventory Management (BOIM)")
        assert "**Inventory Management (BOIM)**" in result


# ---------------------------------------------------------------------------
# _extract_bold_phrases_for_line
# ---------------------------------------------------------------------------


class TestExtractBoldPhrasesForLine:
    def _span(self, text, bold=True, size=11.0):
        return {"text": text, "bold": bold, "size": size}

    def test_extracts_bold_spans(self):
        # Multi-word bold phrase captured; single-word "Teck" is not (too generic when alone)
        spans = [self._span("Data Scientist"), self._span(" at ", bold=False), self._span("Teck")]
        phrases = _extract_bold_phrases_for_line(spans, body_size=11.0)
        assert "Data Scientist" in phrases
        assert "Teck" not in phrases

    def test_skips_non_bold_spans(self):
        spans = [self._span("plain", bold=False)]
        assert _extract_bold_phrases_for_line(spans, body_size=11.0) == []

    def test_skips_header_sized_spans(self):
        # Header spans (size much larger than body) should be skipped
        spans = [self._span("EXPERIENCE", bold=True, size=14.0)]
        assert _extract_bold_phrases_for_line(spans, body_size=11.0) == []

    def test_skips_known_sections(self):
        spans = [self._span("experience")]
        assert _extract_bold_phrases_for_line(spans, body_size=11.0) == []

    def test_skips_single_char(self):
        spans = [self._span("A")]
        assert _extract_bold_phrases_for_line(spans, body_size=11.0) == []

    def test_deduplicates(self):
        # Multi-word bold phrase appearing twice in a mixed line is deduplicated
        spans = [
            self._span("Teck Resources"),
            {"text": " | other", "bold": False, "size": 11.0},
            self._span("Teck Resources"),
        ]
        phrases = _extract_bold_phrases_for_line(spans, body_size=11.0)
        assert phrases.count("Teck Resources") == 1

    def test_sorted_longest_first(self):
        spans = [self._span("T"), self._span("Teck Resources")]
        # "T" is < 2 chars and filtered; "Teck Resources" survives
        phrases = _extract_bold_phrases_for_line(spans, body_size=11.0)
        assert phrases == ["Teck Resources"]

    def test_stop_words_filtered(self):
        # Common stop words like "and" must not be captured as bold phrases
        spans = [self._span("and")]
        phrases = _extract_bold_phrases_for_line(spans, body_size=11.0)
        assert "and" not in phrases


# ---------------------------------------------------------------------------
# _extract_bold_phrases_per_line
# ---------------------------------------------------------------------------


class TestExtractBoldPhrasesPerLine:
    def _span(self, text, bold=True, size=11.0):
        return {"text": text, "bold": bold, "size": size}

    def test_groups_by_line(self):
        lines = [
            [self._span("Data Scientist")],
            [self._span("plain", bold=False)],
            [self._span("Teck Resources")],
        ]
        result = _extract_bold_phrases_per_line(lines, body_size=11.0)
        assert len(result) == 2  # only lines with bold phrases
        assert result[0][0] == "Data Scientist"
        assert result[0][1] == "Data Scientist"
        assert result[1][0] == "Teck Resources"

    def test_empty_lines(self):
        assert _extract_bold_phrases_per_line([], body_size=11.0) == []

    def test_no_bold_spans(self):
        lines = [[self._span("plain", bold=False)]]
        assert _extract_bold_phrases_per_line(lines, body_size=11.0) == []


# ---------------------------------------------------------------------------
# apply_original_bold
# ---------------------------------------------------------------------------


BOLD_PHRASES = [
    ["Data Scientist, 08/2020 - 09/2025", "Data Scientist"],
    ["Teck Resources, Vancouver, BC", "Teck Resources"],
    [
        "Blend Optimization and Inventory Management (BOIM)",
        "Blend Optimization",
        "and",
        "Inventory Management (BOIM)",
    ],
]


class TestApplyOriginalBold:
    def test_no_phrases_passthrough(self):
        text = "Jane Doe\nWork Experience\n- Did things"
        assert apply_original_bold(text, []) == text

    def test_name_line_never_bolded(self):
        text = "Teck Resources\nTeck Resources, Vancouver"
        result = apply_original_bold(text, BOLD_PHRASES)
        lines = result.splitlines()
        assert "**" not in lines[0]

    def test_section_header_never_bolded(self):
        text = "Jane Doe\nWORK EXPERIENCE\nTeck Resources, Vancouver, BC"
        result = apply_original_bold(text, BOLD_PHRASES)
        lines = result.splitlines()
        assert "**" not in lines[1]

    def test_matching_line_gets_bold(self):
        text = "Jane Doe\nData Scientist, 08/2020 - 09/2025"
        result = apply_original_bold(text, BOLD_PHRASES)
        assert "**Data Scientist**" in result

    def test_non_matching_line_unchanged(self):
        text = "Jane Doe\nSoftware Engineer, 2019 - 2021"
        result = apply_original_bold(text, BOLD_PHRASES)
        # Jaccard < 0.5 vs all original lines — no bold applied
        assert "**" not in result.splitlines()[1]

    def test_legacy_flat_format(self):
        # Flat list of strings (legacy format) still works
        text = "Jane Doe\nTeck Resources, Vancouver"
        result = apply_original_bold(text, ["Teck Resources"])
        assert "**Teck Resources**" in result

    def test_phrase_with_parentheses_bolded(self):
        text = "Jane Doe\nBlend Optimization and Inventory Management (BOIM) project"
        result = apply_original_bold(text, BOLD_PHRASES)
        assert "**Inventory Management (BOIM)**" in result

    def test_blank_lines_preserved(self):
        text = "Jane Doe\n\nTeck Resources, Vancouver, BC"
        result = apply_original_bold(text, BOLD_PHRASES)
        assert "\n\n" in result


# ---------------------------------------------------------------------------
# parse_resume_text
# ---------------------------------------------------------------------------


MINIMAL_RESUME = """\
Jane Doe
Vancouver, BC\t416-527-2903
jane@example.com

Data Scientist

WORK EXPERIENCE

Teck Resources, Vancouver, BC
Data Scientist, 08/2020 - 09/2025
- Developed scalable data pipelines
- Improved model accuracy by 30%

EDUCATION

University of Toronto
Bachelor of Science, Computer Science
"""


class TestParseResumeText:
    def _profile(self, **kwargs):
        return StyleProfile(**kwargs)

    def test_first_line_is_name(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        assert els[0].kind == "name"
        assert els[0].text == "Jane Doe"

    def test_contact_lines_classified(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        contacts = [e for e in els if e.kind == "contact"]
        all_texts = [c.text + " " + c.right_text for c in contacts]
        assert any("416" in t for t in all_texts)
        assert any("jane@example.com" in t for t in all_texts)

    def test_tagline_classified_as_tagline(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        taglines = [e for e in els if e.kind == "tagline"]
        assert any("Data Scientist" in t.text for t in taglines)

    def test_known_section_classified_as_section_header(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        headers = [e for e in els if e.kind == "section_header"]
        header_texts = [h.text for h in headers]
        assert "WORK EXPERIENCE" in header_texts
        assert "EDUCATION" in header_texts

    def test_bullet_lines_classified(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        bullets = [e for e in els if e.kind == "bullet"]
        assert len(bullets) == 2
        assert all("Developed" in b.text or "Improved" in b.text for b in bullets)

    def test_body_lines_classified(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        bodies = [e for e in els if e.kind == "body"]
        body_texts = [b.text for b in bodies]
        assert any("Teck Resources" in t for t in body_texts)

    def test_blank_lines_classified(self):
        els = parse_resume_text(MINIMAL_RESUME, self._profile())
        blanks = [e for e in els if e.kind == "blank"]
        assert len(blanks) > 0

    def test_empty_text_returns_empty(self):
        assert parse_resume_text("", self._profile()) == []

    def test_whitespace_only_returns_empty(self):
        assert parse_resume_text("   \n  \n  ", self._profile()) == []

    def test_header_all_caps_detects_uppercase_headers(self):
        text = "Jane Doe\njane@example.com\nEXPERIENCE\n- Did things"
        profile = self._profile(header_all_caps=True)
        els = parse_resume_text(text, profile)
        headers = [e for e in els if e.kind == "section_header"]
        assert any("EXPERIENCE" in h.text for h in headers)

    def test_bullet_char_variants(self):
        text = "Jane Doe\njane@example.com\nEXPERIENCE\n• Bullet one\n▪ Bullet two"
        els = parse_resume_text(text, self._profile())
        bullets = [e for e in els if e.kind == "bullet"]
        assert len(bullets) == 2

    def test_dash_bullet_detected(self):
        text = "Jane Doe\njane@example.com\nEXPERIENCE\n- Dash bullet"
        els = parse_resume_text(text, self._profile())
        bullets = [e for e in els if e.kind == "bullet"]
        assert len(bullets) == 1

    def test_long_contact_block_line_exits_contact_zone(self):
        # A line longer than 80 chars in the contact block exits as body
        long_line = "A" * 85
        text = f"Jane Doe\n{long_line}\nEXPERIENCE\n- thing"
        els = parse_resume_text(text, self._profile())
        kinds = [e.kind for e in els]
        assert "body" in kinds


# ---------------------------------------------------------------------------
# generate_formatted_pdf – smoke test
# ---------------------------------------------------------------------------


class TestGenerateFormattedPdf:
    def test_returns_valid_pdf_bytes(self):
        from career_copilot.resume_formatter.pdf_builder import generate_formatted_pdf

        text = MINIMAL_RESUME
        profile = StyleProfile()
        pdf_bytes = generate_formatted_pdf(text, profile)
        assert isinstance(pdf_bytes, bytes)
        assert pdf_bytes.startswith(b"%PDF")

    def test_centered_resume(self):
        from career_copilot.resume_formatter.pdf_builder import generate_formatted_pdf

        profile = StyleProfile(name_align="center")
        pdf_bytes = generate_formatted_pdf(MINIMAL_RESUME, profile)
        assert pdf_bytes.startswith(b"%PDF")

    def test_with_section_rule(self):
        from career_copilot.resume_formatter.pdf_builder import generate_formatted_pdf

        profile = StyleProfile(
            has_section_rule=True,
            section_rule_color="#1f497d",
            section_rule_thickness=2.25,
            section_rule_style="thickThinSmallGap",
        )
        pdf_bytes = generate_formatted_pdf(MINIMAL_RESUME, profile)
        assert pdf_bytes.startswith(b"%PDF")

    def test_with_bold_phrases(self):
        from career_copilot.resume_formatter.pdf_builder import generate_formatted_pdf

        text = "Jane Doe\njane@example.com\nWORK EXPERIENCE\n**Teck Resources**, Vancouver"
        profile = StyleProfile()
        pdf_bytes = generate_formatted_pdf(text, profile)
        assert pdf_bytes.startswith(b"%PDF")


# ---------------------------------------------------------------------------
# generate_formatted_docx – smoke test
# ---------------------------------------------------------------------------


class TestGenerateFormattedDocx:
    def test_returns_valid_docx_bytes(self):
        from career_copilot.resume_formatter.docx_builder import generate_formatted_docx

        docx_bytes = generate_formatted_docx(MINIMAL_RESUME, StyleProfile())
        assert isinstance(docx_bytes, bytes)
        # DOCX is a zip archive starting with PK
        assert docx_bytes[:2] == b"PK"

    def test_centered_resume(self):
        from career_copilot.resume_formatter.docx_builder import generate_formatted_docx

        profile = StyleProfile(name_align="center")
        docx_bytes = generate_formatted_docx(MINIMAL_RESUME, profile)
        assert docx_bytes[:2] == b"PK"

    def test_with_section_rule(self):
        from career_copilot.resume_formatter.docx_builder import generate_formatted_docx

        profile = StyleProfile(
            has_section_rule=True,
            section_rule_color="#1f497d",
            section_rule_thickness=2.25,
            section_rule_style="thickThinSmallGap",
        )
        docx_bytes = generate_formatted_docx(MINIMAL_RESUME, profile)
        assert docx_bytes[:2] == b"PK"

    def test_docx_contains_name(self):
        from docx import Document

        from career_copilot.resume_formatter.docx_builder import generate_formatted_docx

        docx_bytes = generate_formatted_docx(MINIMAL_RESUME, StyleProfile())
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Doe" in all_text

    def test_docx_contains_section_headers(self):
        from docx import Document

        from career_copilot.resume_formatter.docx_builder import generate_formatted_docx

        docx_bytes = generate_formatted_docx(MINIMAL_RESUME, StyleProfile())
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "WORK EXPERIENCE" in all_text
        assert "EDUCATION" in all_text
