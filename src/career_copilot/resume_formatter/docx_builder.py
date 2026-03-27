"""
Generate a formatted Word (.docx) document from improved resume text
using the original's StyleProfile.
"""

from __future__ import annotations

from io import BytesIO

from career_copilot.resume_formatter.structure_parser import (
    StyleProfile,
    parse_resume_text,
    split_inline_bold,
)


def _docx_font(family: str) -> str:
    """Map a reportlab-style font family to a Word font name."""
    mapping = {
        "Helvetica": "Calibri",
        "Times-Roman": "Times New Roman",
        "Courier": "Courier New",
    }
    return mapping.get(family, "Calibri")


def _pt_to_cm(pt_val: float) -> float:
    return pt_val / 28.35


def generate_formatted_docx(improved_text: str, profile: StyleProfile) -> bytes:
    """
    Render the improved resume text into a Word document that clones the visual
    style described by *profile*.  Returns raw .docx bytes.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Cm, Pt, RGBColor

    def _hex_to_rgb(hex_color: str) -> RGBColor:
        h = hex_color.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    doc = Document()

    # Reset Normal style so it doesn't bleed unexpected defaults onto our paragraphs
    font_name = profile.raw_font_name or _docx_font(profile.base_font_family)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(profile.body_font_size)
        normal.font.bold = False
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.left_indent = Pt(0)
        normal.paragraph_format.first_line_indent = Pt(0)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = profile.line_spacing
    except Exception:
        pass

    # Page margins
    margin_cm = Cm(max(_pt_to_cm(profile.margin_left), 1.27))
    top_cm = Cm(max(_pt_to_cm(profile.margin_top), 1.27))
    for section in doc.sections:
        section.left_margin = margin_cm
        section.right_margin = margin_cm
        section.top_margin = top_cm
        section.bottom_margin = Cm(1.77)

    center = WD_ALIGN_PARAGRAPH.CENTER
    left = WD_ALIGN_PARAGRAPH.LEFT
    name_alignment = center if profile.name_align == "center" else left

    name_color_rgb = _hex_to_rgb(profile.name_color)
    header_color_rgb = _hex_to_rgb(profile.header_color)
    body_color_rgb = _hex_to_rgb(profile.body_color)

    def _new_para(alignment=left, space_before: float = 0.0, space_after: float = 2.0,
                  line_spacing: float | None = None) -> object:
        """Add a Normal-style paragraph with explicit spacing."""
        p = doc.add_paragraph(style="Normal")
        p.alignment = alignment
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        ls = line_spacing if line_spacing is not None else profile.line_spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = ls
        return p

    def add_bottom_border(para) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        sz = max(1, round(profile.section_rule_thickness * 8))
        raw_color = profile.section_rule_color.lstrip("#") if profile.section_rule_color not in ("#000000", "") else "auto"
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), profile.section_rule_style or "single")
        bottom.set(qn("w:sz"), str(sz))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), raw_color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_runs(p, text: str, font_size: float, color_rgb, base_bold: bool = False) -> None:
        """Add one or more runs to *p*, honouring **inline bold** markers."""
        for segment, is_bold in split_inline_bold(text):
            run = p.add_run(segment)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = base_bold or is_bold
            run.font.color.rgb = color_rgb

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    # A4 page width in pt = 595.28; compute usable text width
    _page_width_pt = 595.28
    _text_width_pt = _page_width_pt - 2.0 * max(profile.margin_left, 36.0)
    _text_width_twips = int(_text_width_pt * 20)

    def _add_right_tab(para) -> None:
        """Add a right-aligned tab stop at the right edge of the text area."""
        pPr = para._p.get_or_add_pPr()
        tabs_el = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(_qn("w:val"), "right")
        tab.set(_qn("w:pos"), str(_text_width_twips))
        tabs_el.append(tab)
        pPr.append(tabs_el)

    elements = parse_resume_text(improved_text, profile)

    for el in elements:
        if el.kind == "blank":
            p = _new_para(space_after=2.0)

        elif el.kind == "name":
            p = _new_para(alignment=name_alignment, space_after=profile.name_space_after,
                          line_spacing=1.0)
            run = p.add_run(el.text)
            run.font.name = font_name
            run.font.size = Pt(profile.name_font_size)
            run.font.bold = profile.name_bold
            run.font.color.rgb = name_color_rgb

        elif el.kind == "contact":
            # Split contacts (left + right) must use LEFT alignment so the tab stop works correctly
            _contact_align = left if el.right_text else name_alignment
            p = _new_para(alignment=_contact_align, space_after=profile.contact_space_after)
            if el.right_text:
                _add_right_tab(p)
                add_runs(p, el.text, profile.contact_font_size, body_color_rgb)
                run = p.add_run("\t" + el.right_text)
                run.font.name = font_name
                run.font.size = Pt(profile.contact_font_size)
                run.font.color.rgb = body_color_rgb
            else:
                add_runs(p, el.text, profile.contact_font_size, body_color_rgb)

        elif el.kind == "tagline":
            p = _new_para(alignment=name_alignment, space_after=profile.tagline_space_after)
            add_runs(p, el.text, profile.header_font_size, header_color_rgb,
                     base_bold=profile.header_bold)

        elif el.kind == "section_header":
            p = _new_para(space_before=profile.header_space_before,
                          space_after=profile.header_space_after,
                          line_spacing=1.0)
            run = p.add_run(el.text)
            run.font.name = font_name
            run.font.size = Pt(profile.header_font_size)
            run.font.bold = profile.header_bold
            run.font.color.rgb = header_color_rgb
            if profile.has_section_rule:
                add_bottom_border(p)

        elif el.kind == "bullet":
            p = _new_para(space_after=profile.bullet_space_after)
            p.paragraph_format.left_indent = Pt(14)
            run = p.add_run(f"{profile.bullet_char} ")
            run.font.name = font_name
            run.font.size = Pt(profile.body_font_size)
            run.font.color.rgb = body_color_rgb
            add_runs(p, el.text, profile.body_font_size, body_color_rgb)

        else:  # body
            p = _new_para(space_after=profile.body_space_after)
            if el.right_text:
                _add_right_tab(p)
                add_runs(p, el.text, profile.body_font_size, body_color_rgb)
                run = p.add_run("\t" + el.right_text)
                run.font.name = font_name
                run.font.size = Pt(profile.body_font_size)
                run.font.color.rgb = body_color_rgb
            else:
                add_runs(p, el.text, profile.body_font_size, body_color_rgb)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
