"""
Extract visual style profile from a resume PDF using PyMuPDF.
Also provides shared data classes and plain-text resume parsing
used by the PDF and DOCX builders.
"""

from __future__ import annotations

import json
import re
from typing import Iterator
from dataclasses import asdict, dataclass, field, fields

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

KNOWN_SECTIONS: set[str] = {
    "summary",
    "professional summary",
    "executive summary",
    "objective",
    "profile",
    "about me",
    "experience",
    "work experience",
    "professional experience",
    "employment history",
    "employment",
    "career history",
    "education",
    "academic background",
    "academic history",
    "skills",
    "technical skills",
    "core competencies",
    "key skills",
    "competencies",
    "projects",
    "personal projects",
    "open source",
    "side projects",
    "certifications",
    "certificates",
    "licenses",
    "awards",
    "achievements",
    "honors",
    "languages",
    "volunteer",
    "volunteering",
    "volunteer experience",
    "publications",
    "references",
    "contact",
    "contact information",
    "interests",
    "hobbies",
    "leadership & community involvement",
    "leadership and community involvement",
    "community involvement",
    "professional certifications",
    "professional development",
    "agentic ai & applied machine learning projects",
    "agentic ai projects",
    "applied machine learning projects",
    "independent projects",
    "personal projects & contributions",
    "hackathons",
}

BULLET_CHARS: set[str] = {"•", "·", "○", "▪", "▸", "›", "◦", "▷"}

# Maps PDF/system font names to a reportlab-compatible family string
_FONT_FAMILY_MAP: dict[str, str] = {
    "helvetica": "Helvetica",
    "arial": "Helvetica",
    "calibri": "Helvetica",
    "myriad": "Helvetica",
    "gill": "Helvetica",
    "verdana": "Helvetica",
    "tahoma": "Helvetica",
    "trebuchet": "Helvetica",
    "times": "Times-Roman",
    "georgia": "Times-Roman",
    "garamond": "Times-Roman",
    "palatino": "Times-Roman",
    "cambria": "Times-Roman",
    "courier": "Courier",
}


# ---------------------------------------------------------------------------
# Shared data classes
# ---------------------------------------------------------------------------


@dataclass
class StyleProfile:
    """Captures the visual style of the original resume."""

    name_font_size: float = 18.0
    name_bold: bool = True
    header_font_size: float = 13.0
    header_bold: bool = True
    header_all_caps: bool = False
    body_font_size: float = 11.0
    contact_font_size: float = 10.0
    base_font_family: str = "Helvetica"
    margin_left: float = 50.0
    margin_top: float = 50.0
    sections: list = field(default_factory=list)
    bullet_char: str = "•"
    line_spacing: float = 1.2
    has_header_block: bool = True
    # Alignment: "left" | "center"
    name_align: str = "left"
    # Colors as hex strings e.g. "#2c3e50"
    name_color: str = "#000000"
    header_color: str = "#000000"
    body_color: str = "#000000"
    # Raw font family name from the PDF (e.g. "Calibri") for system font embedding
    raw_font_name: str = ""
    # Bold phrases at body size (e.g. company names, job titles) for post-processing
    bold_phrases: list = field(default_factory=list)
    # Whether the original resume has horizontal rules under section headers
    has_section_rule: bool = False
    # Separator line properties (color hex, thickness in pts, Word border val)
    section_rule_color: str = "#000000"
    section_rule_thickness: float = 0.5
    section_rule_style: str = "single"
    has_name_rule: bool = False
    sections_with_rule: list = field(default_factory=list)
    # Paragraph spacing in points (extracted from original; used by both builders)
    name_space_after: float = 3.0
    contact_space_after: float = 2.0
    header_space_before: float = 10.0
    header_space_after: float = 4.0
    body_space_after: float = 2.0
    bullet_space_after: float = 2.0
    tagline_space_after: float = 2.0

    def to_json(self) -> str:
        return json.dumps(asdict(self))

    @classmethod
    def from_json(cls, s: str) -> "StyleProfile":
        data = json.loads(s)
        valid = {f.name for f in fields(cls)}
        return cls(**{k: v for k, v in data.items() if k in valid})

    @classmethod
    def default(cls) -> "StyleProfile":
        return cls()


_STOP_WORDS: frozenset[str] = frozenset({
    "a", "an", "the", "and", "or", "but", "nor", "so", "yet", "for",
    "in", "on", "at", "to", "of", "by", "as", "is", "was", "are",
    "were", "be", "been", "with", "from", "that", "this", "which",
    "it", "its", "not", "no", "if", "up", "do", "did", "has", "have",
    "had", "may", "can", "will", "all", "both", "each", "per",
})


def _extract_bold_phrases_for_line(line_spans: list[dict], body_size: float) -> list[str]:
    """Return bold phrases for a single line (list of spans), filtered and sorted longest-first."""
    seen: set[str] = set()
    phrases: list[str] = []

    # If ALL non-whitespace spans on this line are bold at body size and long enough,
    # capture the full concatenated line as one phrase (e.g. a standalone job title line).
    # Individual spans are NOT added separately in this case.
    body_spans = [s for s in line_spans if s["text"].strip()]
    whole_line_added = False
    if body_spans and all(
        s.get("bold") and abs(s["size"] - body_size) <= 4.0 and len(s["text"].strip()) >= 2
        for s in body_spans
    ):
        full_line = "".join(s["text"] for s in body_spans).strip()
        lower_full = full_line.lower().rstrip(":")
        if len(full_line) >= 4 and lower_full not in KNOWN_SECTIONS and lower_full not in _STOP_WORDS:
            seen.add(full_line)
            phrases.append(full_line)
            whole_line_added = True

    if not whole_line_added:
        # Mixed line: group consecutive bold body-sized spans into runs.
        # Short bold spans (< 2 chars) are skipped but don't break a run.
        # Only multi-word runs are kept — single-word spans like "Engineer"
        # or "Data" are too generic and get applied in the wrong contexts.
        bold_run: list[str] = []
        for s in line_spans + [None]:  # type: ignore[operator]
            if s is not None and s.get("bold") and abs(s["size"] - body_size) <= 4.0:
                if s["text"].strip() and len(s["text"].strip()) >= 2:
                    bold_run.append(s["text"].strip())
            else:
                if bold_run:
                    phrase = " ".join(bold_run)
                    bold_run = []
                    lower = phrase.lower().rstrip(":")
                    if (len(phrase) >= 2
                            and lower not in KNOWN_SECTIONS
                            and lower not in _STOP_WORDS
                            and len(phrase.split()) >= 2
                            and phrase not in seen):
                        seen.add(phrase)
                        phrases.append(phrase)

    return sorted(phrases, key=len, reverse=True)


def _extract_bold_phrases_per_line(
    lines_data: list[list[dict]], body_size: float
) -> list[list[str]]:
    """
    Return per-line bold phrase map: [[line_text, phrase1, phrase2, ...], ...]
    lines_data: list of lines, each line is a list of span/run dicts with text/size/bold keys.
    Only includes lines that actually have bold phrases.
    """
    result = []
    for line_spans in lines_data:
        phrases = _extract_bold_phrases_for_line(line_spans, body_size)
        if phrases:
            line_text = " ".join(s["text"] for s in line_spans)
            result.append([line_text] + phrases)
    return result


def _apply_phrase_bold(line: str, phrase: str) -> str:
    """Wrap *phrase* in ** within *line*, skipping segments already wrapped."""
    parts = re.split(r"(\*\*.+?\*\*)", line)
    out = []
    for i, part in enumerate(parts):
        if i % 2 == 1:  # already bold
            out.append(part)
        else:
            out.append(re.sub(r"(?<!\w)" + re.escape(phrase) + r"(?!\w)", lambda m: f"**{m.group(0)}**", part, flags=re.IGNORECASE))
    return "".join(out)


def _line_words(s: str) -> frozenset[str]:
    return frozenset(re.findall(r"\w+", s.lower()))


def apply_original_bold(text: str, bold_phrases: list) -> str:
    """
    Post-process LLM-generated resume text: wrap phrases that were bold in the original
    with ** markers wherever they appear.  Skips the name line and section headers.
    """
    if not bold_phrases:
        return text

    # Collect all unique phrases, longest first (avoids nested partial matches)
    seen: set[str] = set()
    all_phrases: list[str] = []
    for entry in bold_phrases:
        phrases = entry[1:] if isinstance(entry, list) and len(entry) >= 2 else (
            [entry] if isinstance(entry, str) else []
        )
        for phrase in phrases:
            if phrase not in seen:
                seen.add(phrase)
                all_phrases.append(phrase)
    all_phrases.sort(key=len, reverse=True)

    lines = text.splitlines()
    result = []
    for i, line in enumerate(lines):
        if i == 0:
            result.append(line)
            continue
        stripped = line.strip().rstrip(":")
        if stripped.lower() in KNOWN_SECTIONS or (
            stripped.isupper() and 1 < len(stripped) < 45
        ):
            result.append(line)
            continue

        if not stripped:
            result.append(line)
            continue

        if stripped[0] in BULLET_CHARS or (len(stripped) > 2 and stripped[0] in "-*" and stripped[1] == " "):
            result.append(line)
            continue

        for phrase in all_phrases:
            line = _apply_phrase_bold(line, phrase)

        result.append(line)
    return "\n".join(result)



def split_inline_bold(text: str) -> list[tuple[str, bool]]:
    """Split a line containing **markers** into (segment, is_bold) pairs.

    Example: 'Worked at **Acme Corp** as **Engineer**'
      → [('Worked at ', False), ('Acme Corp', True), (' as ', False), ('Engineer', True)]
    """
    parts = re.split(r"\*\*(.+?)\*\*", text)
    return [(part, bool(i % 2)) for i, part in enumerate(parts) if part]


@dataclass
class ResumeElement:
    """A single structural element of a plain-text resume."""

    kind: str  # "name" | "contact" | "section_header" | "bullet" | "body" | "blank" | "header_rule"
    text: str
    right_text: str = ""
    has_rule: bool = False  # True if this section_header has an HR rule below it in the original


# ---------------------------------------------------------------------------
# Plain-text resume parser (shared by PDF and DOCX builders)
# ---------------------------------------------------------------------------


_CONTACT_INFO_RE = re.compile(
    r"@|https?://|www\.|linkedin\.|github\.|\.com\b|\b\d{3}[-.\s]\d{3}|\(\d{3}\)"
)

_PHONE_RE = re.compile(r'\b\d{3}[-.\s]\d{3}[-.\s]\d{4}\b')
_EMAIL_RE = re.compile(r'[\w.+%-]+@[\w.-]+\.\w+')
_URL_RE = re.compile(r'(?:linkedin\.|github\.|https?://|www\.)\S+')
_DATE_RANGE_RE = re.compile(
    r'\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\s+\d{4}\s*[–\-]\s*'
    r'(?:(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\s+\d{4}|Present|Current|Now)',
    re.IGNORECASE,
)
_LOCATION_END_RE = re.compile(
    r'\b([A-Z][a-zA-Z\s\-\.\']+,\s*[A-Z]{2,3}(?:,\s*[A-Za-z\s]+)?)\s*$'
)


def _split_contact_line(line: str) -> tuple[str, str] | None:
    """Split 'Location  Phone' or 'Email  URL' into (left, right). Returns None if not splittable."""
    # Phone number at end
    m = _PHONE_RE.search(line)
    if m and m.start() > 5:
        left = line[:m.start()].rstrip()
        right = line[m.start():].strip()
        if left:
            return left, right
    # URL (linkedin/github) after email
    email_m = _EMAIL_RE.search(line)
    url_m = _URL_RE.search(line)
    if email_m and url_m and url_m.start() > email_m.end():
        left = line[:url_m.start()].rstrip()
        right = line[url_m.start():].strip()
        if left:
            return left, right
    return None


def _split_body_line(line: str) -> tuple[str, str] | None:
    """Split 'Company  Location' or 'Title  DateRange' into (left, right). Returns None if not splittable."""
    # Date range at end
    m = _DATE_RANGE_RE.search(line)
    if m and m.start() > 5:
        after = line[m.end():].strip()
        if len(after) <= 3:
            left = line[:m.start()].rstrip()
            right = line[m.start():].strip()
            if left:
                return left, right
    # Location at end (e.g. "Vancouver, BC, Canada" or "Tehran, Iran")
    m = _LOCATION_END_RE.search(line)
    if m and m.start() > 5:
        left = line[:m.start()].rstrip()
        right = m.group(1).strip()
        if left and len(left) > 3:
            return left, right
    return None


def _is_contact_info_line(line: str) -> bool:
    """Return True if the line looks like a contact detail (email, URL, phone, etc.)."""
    return bool(_CONTACT_INFO_RE.search(line))


def parse_resume_text(text: str, profile: StyleProfile) -> list[ResumeElement]:
    """
    Parse a plain-text resume into a list of ResumeElements.

    Detection order:
      1. First non-empty line → name
      2. Lines before the first section header → contact info
      3. Lines matching known section names → section header
      4. Lines starting with bullet chars or "- " / "* " → bullet
      5. Everything else → body
    """
    elements: list[ResumeElement] = []
    lines = text.strip().splitlines()
    if not lines:
        return elements

    section_names = {s.lower().rstrip(":") for s in profile.sections}
    section_names.update(KNOWN_SECTIONS)

    name_found = False
    in_contact_block = True  # true until first section header

    for line in lines:
        stripped = line.strip()

        if not stripped:
            elements.append(ResumeElement("blank", ""))
            continue

        text_lower = stripped.rstrip(":").lower()

        # Section header detection
        is_section = text_lower in section_names or (
            profile.header_all_caps and stripped.isupper() and len(stripped) < 45
        )

        if not name_found:
            elements.append(ResumeElement("name", stripped))
            name_found = True
            continue

        if is_section:
            in_contact_block = False
            header_text = stripped.upper() if profile.header_all_caps else stripped.rstrip(":")
            elements.append(ResumeElement("section_header", header_text))
            continue

        if in_contact_block:
            if _is_contact_info_line(stripped):
                split = _split_contact_line(stripped)
                if split:
                    left, right = split
                    elements.append(ResumeElement("contact", left, right_text=right))
                else:
                    elements.append(ResumeElement("contact", stripped))
            elif len(stripped) <= 80:
                # Short non-contact lines in header zone → title/tagline
                plain = re.sub(r"\*\*", "", stripped)
                elements.append(ResumeElement("tagline", plain))
            else:
                # Long lines signal end of header zone → body
                in_contact_block = False
                elements.append(ResumeElement("body", stripped))
            continue

        # Bullet detection
        if stripped and stripped[0] in BULLET_CHARS:
            elements.append(ResumeElement("bullet", stripped[1:].lstrip()))
            continue

        if len(stripped) > 2 and stripped[0] in "-*" and stripped[1] == " ":
            elements.append(ResumeElement("bullet", stripped[2:]))
            continue

        split = _split_body_line(stripped)
        if split:
            left, right = split
            elements.append(ResumeElement("body", left, right_text=right))
        else:
            elements.append(ResumeElement("body", stripped))

    # Set has_rule on section headers that had rules in the original
    _rule_section_set = {s.lower().rstrip(":") for s in (profile.sections_with_rule or [])}
    if _rule_section_set:
        for el in elements:
            if el.kind == "section_header":
                if el.text.lower().rstrip(":") in _rule_section_set:
                    el.has_rule = True

    # Insert a standalone HR rule right after the name element (for name-block dividers)
    if profile.has_name_rule:
        for i, el in enumerate(elements):
            if el.kind == "name":
                elements.insert(i + 1, ResumeElement("header_rule", ""))
                break

    return elements


# ---------------------------------------------------------------------------
# PDF structure extractor
# ---------------------------------------------------------------------------


def _color_to_hex(color_int: int) -> str:
    """Convert a pymupdf integer color (0xRRGGBB) to a CSS hex string."""
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return f"#{r:02x}{g:02x}{b:02x}"


def _base_font_name(font_name: str) -> str:
    """Strip subset prefix and weight/style suffixes to get the bare family name.

    Examples:
      'BCDEEE+Calibri-Bold' → 'Calibri'
      'Calibri-Bold'        → 'Calibri'
      'ArialMT'             → 'Arial'
      'Arial-BoldMT'        → 'Arial'
    """
    # Drop subset prefix e.g. "BCDEEE+Calibri"
    if "+" in font_name:
        font_name = font_name.split("+")[-1]
    # Take everything before the first dash
    base = font_name.split("-")[0].strip()
    # Strip trailing weight/variant suffixes
    for suffix in ("MT", "PS", "Body", "Regular", "Light", "Medium", "Semibold", "SemBd", "Bd"):
        if base.upper().endswith(suffix.upper()) and len(base) > len(suffix) + 2:
            base = base[: -len(suffix)].strip()
    return base


def _normalize_font_family(font_name: str) -> str:
    lower = font_name.lower()
    for key, mapped in _FONT_FAMILY_MAP.items():
        if key in lower:
            return mapped
    return "Helvetica"


def _is_bold(font_name: str, flags: int) -> bool:
    return bool(flags & 16) or "bold" in font_name.lower()


def _find_body_size(spans: list[dict]) -> float:
    """Return the most common font size (body text)."""
    from collections import Counter

    rounded = [round(s["size"] * 2) / 2 for s in spans]
    if not rounded:
        return 11.0
    return Counter(rounded).most_common(1)[0][0]


def _find_header_size(spans: list[dict], name_size: float, body_size: float) -> float:
    """Return the font size used for section headers (between name and body)."""
    sizes = sorted({round(s["size"] * 2) / 2 for s in spans}, reverse=True)
    candidates = [s for s in sizes if body_size < s < name_size - 0.5]
    if candidates:
        return candidates[-1]
    # Headers may be same size as body but bold — return body size
    return body_size


def _extract_sections(
    spans: list[dict], header_size: float, body_size: float
) -> list[str]:
    sections: list[str] = []
    for span in spans:
        text_lower = span["text"].strip().rstrip(":").lower()
        is_header_sized = abs(span["size"] - header_size) < 1.0 or (
            span["bold"] and abs(span["size"] - body_size) < 1.0
        )
        if is_header_sized and text_lower in KNOWN_SECTIONS:
            display = span["text"].strip().rstrip(":")
            if display not in sections:
                sections.append(display)
    return sections


def _extract_pdf_spacing(
    spans: list[dict], name_size: float, header_size: float, body_size: float
) -> dict[str, float]:
    """Analyse y-coordinate gaps between span rows to estimate paragraph spacing."""
    if not spans:
        return {}

    # Group spans into visual rows by y-proximity (within 2pt = same line)
    rows: list[dict] = []
    for span in spans:
        bbox = span["bbox"]
        y_top, y_bot, sz = float(bbox[1]), float(bbox[3]), float(span["size"])
        if rows and abs(y_top - rows[-1]["y_top"]) < 2.0:
            rows[-1]["y_bot"] = max(rows[-1]["y_bot"], y_bot)
            rows[-1]["size"] = max(rows[-1]["size"], sz)
        else:
            rows.append({"y_top": y_top, "y_bot": y_bot, "size": sz})

    if len(rows) < 2:
        return {}

    def _role(size: float) -> str:
        if abs(size - name_size) < 1.0:
            return "name"
        if abs(size - header_size) < 1.0:
            return "header"
        if size < body_size - 0.5:
            return "contact"
        return "body"

    gaps_after: dict[str, list[float]] = {"name": [], "contact": [], "header": [], "body": []}
    gaps_before_header: list[float] = []
    for i in range(len(rows) - 1):
        gap = max(0.0, rows[i + 1]["y_top"] - rows[i]["y_bot"])
        role = _role(rows[i]["size"])
        gaps_after[role].append(gap)
        if _role(rows[i + 1]["size"]) == "header":
            gaps_before_header.append(gap)

    result = {}
    for role, gaps in gaps_after.items():
        if gaps:
            s = sorted(gaps)
            median = s[len(s) // 2]
            result[f"{role}_gap"] = round(median, 1)
    if gaps_before_header:
        s = sorted(gaps_before_header)
        result["header_before_gap"] = round(s[len(s) // 2], 1)
    return result


def _compute_pdf_line_spacing(spans: list[dict], body_size: float) -> float:
    """Estimate line spacing multiplier from the y-distance between tops of consecutive body rows."""
    body_rows = sorted(
        [s for s in spans if abs(s["size"] - body_size) < 1.0],
        key=lambda s: float(s["bbox"][1]),
    )
    if len(body_rows) < 2:
        return 1.2
    leadings: list[float] = []
    for i in range(len(body_rows) - 1):
        top_gap = float(body_rows[i + 1]["bbox"][1]) - float(body_rows[i]["bbox"][1])
        # Only count gaps that look like in-paragraph leading (not paragraph breaks)
        if body_size * 0.9 < top_gap < body_size * 2.2:
            leadings.append(top_gap / body_size)
    if not leadings:
        return 1.2
    s = sorted(leadings)
    return round(s[len(s) // 2], 2)


def parse_resume_structure(pdf_bytes: bytes) -> StyleProfile:
    """
    Analyse a resume PDF and return a StyleProfile capturing its visual style.
    Falls back to StyleProfile.default() if the PDF has no extractable text.
    """
    try:
        import pymupdf
    except ImportError:
        return StyleProfile.default()

    try:
        doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return StyleProfile.default()

    spans: list[dict] = []
    pdf_lines_data: list[list[dict]] = []

    # Use page 0 for layout metrics and drawing detection
    page = doc[0]
    try:
        _p0_blocks = page.get_text("dict")["blocks"]
    except Exception:
        doc.close()
        return StyleProfile.default()

    def _extract_page_spans(blocks_list):
        for block in blocks_list:
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_span_list: list[dict] = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    entry = {
                        "text": text,
                        "size": float(span.get("size", 11.0)),
                        "font": span.get("font", "Helvetica"),
                        "flags": int(span.get("flags", 0)),
                        "bold": _is_bold(
                            span.get("font", ""), int(span.get("flags", 0))
                        ),
                        "bbox": span.get("bbox", (50.0, 50.0, 500.0, 60.0)),
                        "color": int(span.get("color", 0)),
                    }
                    spans.append(entry)
                    line_span_list.append(entry)
                if line_span_list:
                    pdf_lines_data.append(line_span_list)

    _extract_page_spans(_p0_blocks)
    _page_span_ends: list[int] = [len(spans)]

    # Extract spans from remaining pages (needed for bold phrase detection across all pages)
    for pg_idx in range(1, len(doc)):
        try:
            _extra_blocks = doc[pg_idx].get_text("dict")["blocks"]
            _extract_page_spans(_extra_blocks)
        except Exception:
            pass
        _page_span_ends.append(len(spans))

    page_width = page.rect.width

    # Snapshot drawings from ALL pages before closing (page becomes invalid after close)
    _all_page_drawings: list[tuple[int, list]] = []
    for _pg_i in range(len(doc)):
        try:
            _all_page_drawings.append((_pg_i, list(doc[_pg_i].get_drawings())))
        except Exception:
            _all_page_drawings.append((_pg_i, []))

    doc.close()

    if not spans:
        return StyleProfile.default()

    # Font size roles
    sizes = sorted({s["size"] for s in spans}, reverse=True)
    name_size = sizes[0]
    body_size = _find_body_size(spans)
    header_size = _find_header_size(spans, name_size, body_size)
    contact_size = min(sizes) if len(sizes) > 2 else max(body_size - 1.0, 8.0)

    # Base font family (from most common body-sized spans)
    body_spans = [s for s in spans if abs(s["size"] - body_size) < 1.0]
    base_font = "Helvetica"
    raw_font = ""
    if body_spans:
        raw_font = _base_font_name(body_spans[0]["font"])
        base_font = _normalize_font_family(body_spans[0]["font"])

    # Name styling
    name_spans = [s for s in spans if abs(s["size"] - name_size) < 0.5]
    name_bold = any(s["bold"] for s in name_spans) if name_spans else True

    # Header styling
    header_spans = [s for s in spans if abs(s["size"] - header_size) < 0.5]
    header_bold = any(s["bold"] for s in header_spans) if header_spans else True
    header_all_caps = (
        bool(header_spans)
        and all(s["text"].isupper() for s in header_spans[:6] if len(s["text"]) > 2)
    )

    # Sections in order
    sections = _extract_sections(spans, header_size, body_size)

    # Bullet character
    bullet_char = "•"
    for s in spans:
        ch = s["text"].lstrip()[:1]
        if ch in BULLET_CHARS:
            bullet_char = ch
            break

    # Margins
    lefts = [s["bbox"][0] for s in spans]
    tops = [s["bbox"][1] for s in spans]
    margin_left = round(min(lefts), 1) if lefts else 50.0
    margin_top = round(min(tops), 1) if tops else 50.0

    # Has name/contact header block
    has_header_block = bool(name_spans)

    # Name alignment: use the combined bounding box of all name spans
    # (the name may be split across multiple spans, so one span's center is unreliable)
    name_align = "left"
    if name_spans:
        combined_left = min(s["bbox"][0] for s in name_spans)
        combined_right = max(s["bbox"][2] for s in name_spans)
        text_center_x = (combined_left + combined_right) / 2
        if abs(text_center_x - page_width / 2) < page_width * 0.15:
            name_align = "center"

    # Colors (most common color per role)
    from collections import Counter

    def _dominant_color(span_list: list[dict]) -> str:
        if not span_list:
            return "#000000"
        counts = Counter(s["color"] for s in span_list)
        return _color_to_hex(counts.most_common(1)[0][0])

    name_color = _dominant_color(name_spans)
    header_color = _dominant_color(header_spans)
    body_color = _dominant_color(body_spans)

    # Paragraph spacing from y-coordinate gap analysis
    _pdf_gaps = _extract_pdf_spacing(spans, name_size, header_size, body_size)
    def _clamped(key: str, default: float) -> float:
        v = _pdf_gaps.get(key)
        return round(v, 1) if v is not None and 0 <= v <= 30 else default
    pdf_name_space_after    = _clamped("name_gap",         3.0)
    pdf_contact_space_after = _clamped("contact_gap",      2.0)
    pdf_header_space_before = _clamped("header_before_gap", 10.0)
    pdf_header_space_after  = _clamped("header_gap",        4.0)
    pdf_body_space_after    = _clamped("body_gap",          2.0)
    pdf_bullet_space_after  = pdf_body_space_after
    pdf_tagline_space_after = _clamped("name_gap",          2.0)
    pdf_line_spacing        = _compute_pdf_line_spacing(spans, body_size)

    # Detect horizontal separator lines per-page, matching each to a section header or name block
    pdf_has_section_rule = False
    pdf_rule_color = header_color
    pdf_rule_thickness = 0.5
    _has_name_rule = False
    _sections_with_rule_pdf: list[str] = []
    _rule_ys_for_style: list[tuple[float, float]] = []  # (y, thickness) for style classification

    def _pg_spans(pg_idx: int) -> list[dict]:
        start = _page_span_ends[pg_idx - 1] if pg_idx > 0 else 0
        end = _page_span_ends[pg_idx] if pg_idx < len(_page_span_ends) else len(spans)
        return spans[start:end]

    def _header_positions(page_spans: list[dict]) -> list[tuple[float, str]]:
        result: list[tuple[float, str]] = []
        for s in page_spans:
            is_hdr = abs(s["size"] - header_size) < 1.5
            is_bold_body = s.get("bold") and abs(s["size"] - body_size) < 1.0
            if is_hdr or is_bold_body:
                t = s["text"].strip().rstrip(":").lower()
                if t in KNOWN_SECTIONS:
                    y = (float(s["bbox"][1]) + float(s["bbox"][3])) / 2
                    result.append((y, t))
        return result

    _p0_hdrs = _header_positions(_pg_spans(0))
    _min_section_y_p0 = min((y for y, _ in _p0_hdrs), default=None)

    try:
        for pg_idx, pg_drawings in _all_page_drawings:
            pg_hdrs = _header_positions(_pg_spans(pg_idx))

            for drawing in pg_drawings:
                found_rule = False
                rule_y = 0.0
                rule_thick = 0.5
                rule_color_str = pdf_rule_color

                for item in drawing.get("items", []):
                    if item[0] == "l":
                        p1, p2 = item[1], item[2]
                        x0 = p1.x if hasattr(p1, "x") else p1[0]
                        y0 = p1.y if hasattr(p1, "y") else p1[1]
                        x1 = p2.x if hasattr(p2, "x") else p2[0]
                        y1 = p2.y if hasattr(p2, "y") else p2[1]
                        if abs(y1 - y0) < 3 and abs(x1 - x0) > page_width * 0.3:
                            found_rule = True
                            rule_y = (y0 + y1) / 2
                            rule_thick = float(drawing.get("width") or 0.5)
                    elif item[0] == "re":
                        rect = item[1]
                        rw = rect.width if hasattr(rect, "width") else (rect[2] - rect[0])
                        rh = rect.height if hasattr(rect, "height") else (rect[3] - rect[1])
                        ry0 = rect.y0 if hasattr(rect, "y0") else rect[1]
                        if rh < 4 and rw > page_width * 0.3:
                            found_rule = True
                            rule_y = ry0 + rh / 2
                            rule_thick = max(rh, 0.5)
                    if found_rule:
                        stroke = drawing.get("color") or drawing.get("fill")
                        if stroke and len(stroke) >= 3:
                            r2, g2, b2 = int(stroke[0]*255), int(stroke[1]*255), int(stroke[2]*255)
                            rule_color_str = f"#{r2:02x}{g2:02x}{b2:02x}"
                        break

                if not found_rule:
                    continue

                pdf_has_section_rule = True
                _rule_ys_for_style.append((rule_y, rule_thick))
                if pdf_rule_color == header_color:
                    pdf_rule_color = rule_color_str
                    pdf_rule_thickness = rule_thick

                # Name-block rule: page 0, y is above the first section header
                if (pg_idx == 0 and _min_section_y_p0 is not None
                        and rule_y < _min_section_y_p0 - 5):
                    _has_name_rule = True
                    continue

                # Match to the closest section header above this rule (within 35pt)
                best: tuple[float, str] | None = None
                for hdr_y, hdr_name in pg_hdrs:
                    gap = rule_y - hdr_y
                    if 0 < gap <= 35:
                        if best is None or gap < (rule_y - best[0]):
                            best = (hdr_y, hdr_name)
                if best:
                    if best[1] not in _sections_with_rule_pdf:
                        _sections_with_rule_pdf.append(best[1])

    except Exception as _draw_err:
        print(f"[StyleParser] drawing detection error: {_draw_err!r}")

    # Classify rule style: if two rules land within 6pt of each other, it's a double-line border
    pdf_section_rule_style = "single"
    if len(_rule_ys_for_style) >= 2:
        sorted_ys = sorted(_rule_ys_for_style, key=lambda r: r[0])
        for _si in range(len(sorted_ys) - 1):
            y_a, t_a = sorted_ys[_si]
            y_b, t_b = sorted_ys[_si + 1]
            if y_b - y_a <= 6:
                pdf_section_rule_style = "thickThinSmallGap" if t_a >= t_b else "thinThickSmallGap"
                # Use the combined thickness
                pdf_rule_thickness = t_a + t_b
                break

    return StyleProfile(
        name_font_size=round(name_size, 1),
        name_bold=name_bold,
        header_font_size=round(header_size, 1),
        header_bold=header_bold,
        header_all_caps=header_all_caps,
        body_font_size=round(body_size, 1),
        contact_font_size=round(contact_size, 1),
        base_font_family=base_font,
        margin_left=margin_left,
        margin_top=margin_top,
        sections=sections,
        bullet_char=bullet_char,
        line_spacing=pdf_line_spacing,
        has_header_block=has_header_block,
        name_align=name_align,
        name_color=name_color,
        header_color=header_color,
        body_color=body_color,
        raw_font_name=raw_font,
        bold_phrases=_extract_bold_phrases_per_line(pdf_lines_data, body_size),
        has_section_rule=pdf_has_section_rule,
        has_name_rule=_has_name_rule,
        sections_with_rule=_sections_with_rule_pdf,
        section_rule_color=pdf_rule_color,
        section_rule_thickness=pdf_rule_thickness,
        section_rule_style=pdf_section_rule_style,
        name_space_after=pdf_name_space_after,
        contact_space_after=pdf_contact_space_after,
        header_space_before=pdf_header_space_before,
        header_space_after=pdf_header_space_after,
        body_space_after=pdf_body_space_after,
        bullet_space_after=pdf_bullet_space_after,
        tagline_space_after=pdf_tagline_space_after,
    )


def _eff_space(para, attr: str) -> float | None:
    """Return space_before or space_after in pt for a paragraph, following the style chain."""
    try:
        val = getattr(para.paragraph_format, attr, None)
        if val is not None:
            return val.pt
        style = para.style
        while style:
            val = getattr(style.paragraph_format, attr, None)
            if val is not None:
                return val.pt
            style = style.base_style
    except Exception:
        pass
    return None


def parse_resume_structure_docx(docx_bytes: bytes) -> StyleProfile:
    """
    Analyse a resume .docx file and return a StyleProfile capturing its visual style.
    Falls back to StyleProfile.default() if the file cannot be read.
    """
    try:
        from collections import Counter
        from io import BytesIO

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(BytesIO(docx_bytes))
    except Exception:
        return StyleProfile.default()

    # --- Helpers that walk the style inheritance chain ---

    def _eff_size(run, para) -> float:
        if run.font.size:
            return run.font.size.pt
        style = para.style
        while style:
            if style.font.size:
                return style.font.size.pt
            style = style.base_style
        return 11.0

    def _eff_bold(run, para) -> bool:
        if run.bold is not None:
            return bool(run.bold)
        style = para.style
        while style:
            if style.font.bold is not None:
                return bool(style.font.bold)
            style = style.base_style
        return False

    def _eff_font_name(run, para) -> str:
        if run.font.name:
            return run.font.name
        style = para.style
        while style:
            if style.font.name:
                return style.font.name
            style = style.base_style
        return "Calibri"

    def _eff_color(run) -> str:
        # Try run-level XML first (handles theme colors that python-docx can't resolve)
        try:
            from docx.oxml.ns import qn as _qn
            rPr = run._r.find(_qn("w:rPr"))
            if rPr is not None:
                color_el = rPr.find(_qn("w:color"))
                if color_el is not None:
                    val = color_el.get(_qn("w:val"))
                    if val and val.upper() != "AUTO" and len(val) == 6:
                        return f"#{val.lower()}"
        except Exception:
            pass
        # Fallback: python-docx resolved color
        try:
            if run.font.color and run.font.color.type is not None:
                rgb = run.font.color.rgb
                return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
        except Exception:
            pass
        return "#000000"

    # --- Collect all paragraphs (body + table cells) ---

    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    # --- Build run-level data ---

    runs_data: list[dict] = []
    para_meta: list[dict] = []  # one entry per paragraph for alignment lookup

    def _bottom_border_props(para) -> dict | None:
        """Return {"color": "#rrggbb", "thickness": pt, "style": str} if para has a bottom
        border, else None.  Checks inline XML first, then walks the style chain."""
        from docx.oxml.ns import qn as _qn2

        def _parse_border_el(bottom) -> dict | None:
            val = bottom.get(_qn2("w:val"), "none")
            if val in ("none", ""):
                return None
            sz = int(bottom.get(_qn2("w:sz"), "4") or "4")
            thickness = sz / 8.0
            raw_color = bottom.get(_qn2("w:color"), "auto")
            color = f"#{raw_color.lower()}" if raw_color not in ("auto", "") and len(raw_color) == 6 else None
            return {"color": color, "thickness": thickness, "style": val}

        try:
            # 1. Inline paragraph properties
            pPr = para._p.find(_qn2("w:pPr"))
            if pPr is not None:
                pBdr = pPr.find(_qn2("w:pBdr"))
                if pBdr is not None:
                    bottom = pBdr.find(_qn2("w:bottom"))
                    if bottom is not None:
                        result = _parse_border_el(bottom)
                        if result:
                            return result

            # 2. Style chain — the border may be defined in a named style
            style = para.style
            while style:
                el = style.element
                pPr_s = el.find(_qn2("w:pPr"))
                if pPr_s is not None:
                    pBdr_s = pPr_s.find(_qn2("w:pBdr"))
                    if pBdr_s is not None:
                        bottom = pBdr_s.find(_qn2("w:bottom"))
                        if bottom is not None:
                            result = _parse_border_el(bottom)
                            if result:
                                return result
                style = style.base_style
        except Exception:
            pass
        return None

    for para in all_paras:
        para_runs = []
        for run in para.runs:
            if not run.text.strip():
                continue
            entry = {
                "text": run.text.strip(),
                "size": float(_eff_size(run, para)),
                "bold": _eff_bold(run, para),
                "font": _eff_font_name(run, para),
                "color": _eff_color(run),
            }
            runs_data.append(entry)
            para_runs.append(entry)
        if para_runs:
            para_meta.append({"runs": para_runs, "align": para.alignment, "border": _bottom_border_props(para)})

    if not runs_data:
        return StyleProfile.default()

    # --- Font size roles ---

    from collections import Counter

    sizes = sorted({r["size"] for r in runs_data}, reverse=True)
    name_size = sizes[0]
    size_counts = Counter(round(r["size"] * 2) / 2 for r in runs_data)
    body_size = float(size_counts.most_common(1)[0][0])
    candidates = [s for s in sizes if body_size < s < name_size - 0.5]
    header_size = candidates[-1] if candidates else body_size

    name_runs = [r for r in runs_data if abs(r["size"] - name_size) < 0.5]
    header_runs = [r for r in runs_data if abs(r["size"] - header_size) < 0.5]

    # --- Paragraph spacing extraction ---
    _para_size_pairs: list[tuple] = []
    for para in all_paras:
        for run in para.runs:
            if run.text.strip():
                _para_size_pairs.append((para, float(_eff_size(run, para))))
                break

    def _sample_space(role_size: float, bold_only: bool = False) -> tuple:
        for para, sz in _para_size_pairs:
            if abs(sz - role_size) <= 0.5:
                if bold_only and not any(_eff_bold(r, para) for r in para.runs if r.text.strip()):
                    continue
                return _eff_space(para, "space_before"), _eff_space(para, "space_after")
        return None, None

    _name_sb, _name_sa = _sample_space(name_size)
    # When header_size == body_size, target bold paragraphs specifically
    _hdr_bold_only = abs(header_size - body_size) < 0.5
    _hdr_sb, _hdr_sa = _sample_space(header_size, bold_only=_hdr_bold_only)
    _body_sb, _body_sa = _sample_space(body_size)

    docx_name_space_after    = _name_sa if _name_sa is not None else 3.0
    docx_header_space_before = _hdr_sb  if _hdr_sb  is not None else 10.0
    docx_header_space_after  = _hdr_sa  if _hdr_sa  is not None else 4.0
    docx_body_space_after    = _body_sa if _body_sa is not None else 2.0
    docx_contact_space_after = docx_body_space_after
    docx_bullet_space_after  = docx_body_space_after
    docx_tagline_space_after = _name_sa if _name_sa is not None else 2.0

    # Extract line spacing from body paragraphs
    docx_line_spacing = 1.2
    try:
        from docx.enum.text import WD_LINE_SPACING
        for para, sz in _para_size_pairs:
            if abs(sz - body_size) > 0.5:
                continue
            pf = para.paragraph_format
            ls = pf.line_spacing
            lsr = pf.line_spacing_rule
            if ls is None:
                # Walk style chain
                style = para.style
                while style and ls is None:
                    ls = style.paragraph_format.line_spacing
                    lsr = style.paragraph_format.line_spacing_rule
                    style = style.base_style
            if ls is not None:
                if lsr == WD_LINE_SPACING.MULTIPLE:
                    docx_line_spacing = round(float(ls), 2)
                elif hasattr(ls, "pt") and ls.pt and body_size:
                    docx_line_spacing = round(ls.pt / body_size, 2)
                break
    except Exception:
        pass

    section_rule_thickness = 0.5
    section_rule_color_raw = None
    section_rule_style_val = "single"
    _docx_sections_with_rule: list[str] = []
    _docx_has_name_rule = False
    _first_section_seen = False
    _rule_props_first: dict | None = None

    from docx.oxml.ns import qn as _qn3  # noqa: F811
    for pm in para_meta:
        runs_pm = pm["runs"]
        para_text_pm = " ".join(r["text"] for r in runs_pm).strip().rstrip(":")
        text_lower_pm = para_text_pm.lower()
        is_section_pm = text_lower_pm in KNOWN_SECTIONS and (
            any(abs(r["size"] - header_size) < 1.5 for r in runs_pm)
            or any(abs(r["size"] - body_size) < 1.0 and r["bold"] for r in runs_pm)
        )
        if is_section_pm:
            _first_section_seen = True
        if pm.get("border"):
            if _rule_props_first is None:
                _rule_props_first = pm["border"]
            if is_section_pm:
                if text_lower_pm not in _docx_sections_with_rule:
                    _docx_sections_with_rule.append(text_lower_pm)
            elif not _first_section_seen:
                _docx_has_name_rule = True

    has_section_rule = bool(_docx_sections_with_rule) or _docx_has_name_rule
    if _rule_props_first:
        section_rule_thickness = _rule_props_first.get("thickness", 0.5)
        section_rule_color_raw = _rule_props_first.get("color")
        section_rule_style_val = _rule_props_first.get("style", "single")
    body_runs = [r for r in runs_data if abs(r["size"] - body_size) < 1.0]

    name_bold = any(r["bold"] for r in name_runs) if name_runs else True
    header_bold = any(r["bold"] for r in header_runs) if header_runs else True
    header_all_caps = bool(header_runs) and all(
        r["text"].isupper() for r in header_runs[:6] if len(r["text"]) > 2
    )

    # raw_font from body runs (like PDF version), not the name
    raw_font = _base_font_name(body_runs[0]["font"]) if body_runs else "Calibri"
    base_font = _normalize_font_family(raw_font)

    # --- Name alignment: check the paragraph that contains name-sized text ---

    name_align = "left"
    for pm in para_meta:
        if any(abs(r["size"] - name_size) < 0.5 for r in pm["runs"]):
            if pm["align"] == WD_ALIGN_PARAGRAPH.CENTER:
                name_align = "center"
            break

    # --- Colors ---

    def dominant_color(run_list: list[dict]) -> str:
        if not run_list:
            return "#000000"
        counts = Counter(r["color"] for r in run_list)
        return counts.most_common(1)[0][0]

    section_rule_color = section_rule_color_raw or dominant_color(header_runs)
    section_rule_style = section_rule_style_val

    # --- Sections ---

    sections = list(dict.fromkeys(
        r["text"].rstrip(":")
        for r in runs_data
        if abs(r["size"] - header_size) < 0.5
        and r["text"].strip().rstrip(":").lower() in KNOWN_SECTIONS
    ))

    # --- Bullet character ---

    bullet_char = "•"
    for pm in para_meta:
        for r in pm["runs"]:
            ch = r["text"][:1]
            if ch in BULLET_CHARS:
                bullet_char = ch
                break

    # --- Bold body phrases ---

    bold_phrases_list = _extract_bold_phrases_per_line(
        [[{"text": r["text"], "size": r["size"], "bold": r["bold"]} for r in pm["runs"]]
         for pm in para_meta],
        body_size,
    )

    # --- Margins ---

    margin_left = 72.0
    margin_top = 72.0
    try:
        sec = doc.sections[0]
        margin_left = sec.left_margin.pt if sec.left_margin else 72.0
        margin_top = sec.top_margin.pt if sec.top_margin else 72.0
    except Exception:
        pass

    return StyleProfile(
        name_font_size=round(name_size, 1),
        name_bold=name_bold,
        header_font_size=round(header_size, 1),
        header_bold=header_bold,
        header_all_caps=header_all_caps,
        body_font_size=round(body_size, 1),
        contact_font_size=round(min(sizes), 1) if len(sizes) > 2 else max(round(body_size - 1.0, 1), 8.0),
        base_font_family=base_font,
        margin_left=margin_left,
        margin_top=margin_top,
        sections=sections,
        bullet_char=bullet_char,
        line_spacing=docx_line_spacing,
        has_header_block=bool(name_runs),
        name_align=name_align,
        name_color=dominant_color(name_runs),
        header_color=dominant_color(header_runs),
        body_color=dominant_color(body_runs),
        raw_font_name=raw_font,
        bold_phrases=bold_phrases_list,
        has_section_rule=has_section_rule,
        has_name_rule=_docx_has_name_rule,
        sections_with_rule=_docx_sections_with_rule,
        section_rule_color=section_rule_color,
        section_rule_thickness=section_rule_thickness,
        section_rule_style=section_rule_style,
        name_space_after=docx_name_space_after,
        contact_space_after=docx_contact_space_after,
        header_space_before=docx_header_space_before,
        header_space_after=docx_header_space_after,
        body_space_after=docx_body_space_after,
        bullet_space_after=docx_bullet_space_after,
        tagline_space_after=docx_tagline_space_after,
    )
